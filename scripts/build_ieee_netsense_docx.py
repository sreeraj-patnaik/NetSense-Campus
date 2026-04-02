"""
Build IEEE-style technical paper (Word .docx) for NetSense Campus.
Formatting approximates IEEE conference manuscript: Times New Roman 10pt, structured sections,
numbered references, formal abstract and index terms.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "NetSense_IEEE_Paper.docx"


def set_run_font(run, name="Times New Roman", size=Pt(10), bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    run.bold = bold
    run.font.italic = italic


def ieee_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(text)
    set_run_font(r, size=Pt(10), bold=True)
    return p


def body_p(doc, text, first_line_indent=Pt(0)):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_eq_label(doc, equation_text, number):
    """Centered equation line + number at right (simplified as two-line block)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(equation_text)
    set_run_font(r, italic=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run(f"({number})")
    set_run_font(r2)


def add_reference(doc, num, text):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(f"[{num}] ")
    set_run_font(r)
    r2 = p.add_run(text)
    set_run_font(r2)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.625)
    section.right_margin = Inches(0.625)

    # ---- Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(6)
    tr = t.add_run(
        "NetSense Campus: A Web-Native Framework for Grid-Based "
        "Indoor Wi-Fi and Cellular Signal Visualization"
    )
    set_run_font(tr, size=Pt(16), bold=True)

    # ---- Authors
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    a.paragraph_format.space_after = Pt(2)
    ar = a.add_run("Sreeraj Patnaik")
    set_run_font(ar, size=Pt(11))
    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.paragraph_format.space_after = Pt(12)
    affr = aff.add_run("Department of Computer Science and Engineering — (Institution Name), India")
    set_run_font(affr, size=Pt(10), italic=True)

    # ---- Abstract
    ab = doc.add_paragraph()
    ab.paragraph_format.space_after = Pt(6)
    abr = ab.add_run("Abstract—")
    set_run_font(abr, bold=True)
    abst = (
        "Indoor wireless coverage is spatially heterogeneous; operational teams rarely possess a compact, "
        "map-aligned representation that fuses Wi-Fi and cellular received signal strength indicator (RSSI)–class "
        "measurements across carriers and access networks. This paper presents NetSense Campus (NSC), a deployable "
        "web stack that discretizes each building floor into a configurable rectangular grid, ingests samples via "
        "browser or REST API, and maintains per-cell robust aggregates. The system applies the sample median within "
        "each cell for outlier resistance and materializes two aggregate namespaces per mode: provider-specific buckets "
        "and an all-provider composite to avoid runtime scan-table joins on read paths. Empty traversable cells may "
        "be filled by a single-pass, inverse-distance–weighted estimator on the grid with explicit client-side "
        "differentiation of interpolated versus measured values. We specify the data model, ingest validation, "
        "interpolation kernel, heatmap rendering pipeline, and application-layer complexity. NSC is implemented in "
        "Django 5 with SQLite or PostgreSQL backends and is suitable for campus and enterprise indoor planning "
        "workflows where transparency and API simplicity outweigh full physics-based RF simulation."
    )
    ar2 = ab.add_run(abst)
    set_run_font(ar2)

    # ---- Index Terms
    it = doc.add_paragraph()
    it.paragraph_format.space_after = Pt(12)
    itr = it.add_run("Index Terms—")
    set_run_font(itr, bold=True)
    it2 = it.add_run(
        "Indoor wireless networks; received signal strength; heatmap visualization; "
        "median aggregation; inverse distance weighting; campus information systems; web applications."
    )
    set_run_font(it2)

    # ---- I. INTRODUCTION
    ieee_heading(doc, "I. INTRODUCTION")
    body_p(
        doc,
        "Modern universities and enterprises depend on predictable indoor connectivity for learning management "
        "systems, voice, and Internet of Things deployments. End-user complaints are often qualitative "
        "(“slow” or “no bars”) while engineering teams require quantitative, spatially indexed evidence to "
        "place access points, negotiate carrier service, or justify infrastructure spend [1], [2].",
    )
    body_p(
        doc,
        "NetSense Campus addresses this need with a deliberately minimal architecture: each floor plan is "
        "registered with horizontal dimensions (rows and columns), optional raster underlays, and a set of "
        "blocked cell identifiers representing non-traversable regions. Samples are atomic: each observation "
        "binds a mode (Wi-Fi or mobile), optional service-provider label, network identifier, and integer "
        "signal strength in dBm. The contribution is not a new RF propagation law but an integrated, "
        "openly specified pipeline—from ingestion through median-based summarization, optional gap filling, "
        "and color-mapped rendering—that remains tractable for student or small-team maintenance.",
    )

    ieee_heading(doc, "II. RELATED WORK AND PROBLEM FORMULATION")
    body_p(
        doc,
        "Indoor radio mapping spans drive-test cartography, ray-tracing-based planning tools, and crowd-sourced "
        "platforms. Commercial planning suites offer high fidelity at the cost of licensing and calibration "
        "overhead [3]. Academic work on fingerprinting and SLAM-assisted localization emphasizes device position "
        "estimation rather than stakeholder-facing coverage cartography [4]. Spreadsheet-centric site surveys "
        "lack referential integrity between numeric readings and floor geometry.",
    )
    body_p(
        doc,
        "NSC occupies a middle tier: it assumes operators can align a regular grid to a floor image administratively, "
        "then scales to hundreds of cells per floor with O(S) aggregate rebuilds over S raw scans and "
        "O(R·C) interpolation cost for R rows and C columns under a fixed neighborhood radius. The design question "
        "is therefore how to balance statistical robustness, explainability, and API surface area—not "
        "millimeter-wave accuracy.",
    )

    ieee_heading(doc, "III. SYSTEM ARCHITECTURE")
    body_p(
        doc,
        "Figure 1 summarizes control and data flow. Clients POST JSON or form-encoded payloads to /api/scan/; "
        "the server validates block, floor, mode, provider (against configurable allow-lists), and cell "
        "membership, persisting a Scan row and synchronously refreshing CellAggregate tuples for the affected "
        "cell. Read clients request /api/heatmap/ with block, floor, mode, optional provider filter, and an "
        "interpolation toggle; the response is a JSON array consumed by a browser canvas renderer.",
    )
    fig = doc.add_paragraph()
    fig.paragraph_format.space_before = Pt(6)
    fig.paragraph_format.space_after = Pt(6)
    fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fig.add_run(
        "Fig. 1. Dataflow: scan ingestion updates aggregates; heatmap queries optionally append interpolated cells."
    )
    set_run_font(fr, size=Pt(9), italic=True)

    ieee_heading(doc, "IV. DATA MODEL AND GRID FORMALISM")
    body_p(
        doc,
        "Let each active floor plan ω specify grid width N_x and height N_y. A traversable cell is addressed "
        "by coordinates (x, y) with 0 ≤ x < N_x and 0 ≤ y < N_y. For compact storage and blocked-cell "
        "membership checks, NSC uses a row-major linear index:",
    )
    add_eq_label(doc, "c(x, y) = y · N_x + x", 1)
    body_p(
        doc,
        "The inverse mapping recovers x = c mod N_x and y = ⌊c / N_x⌋. Blocked cells form a set B_ω ⊆ "
        "{0, …, N_x N_y − 1}; ingests with (x, y) ⟼ c ∈ B_ω are rejected. This indexing is mirrored in the "
        "JavaScript scan UI so administrative configuration and client behavior remain consistent.",
    )
    body_p(
        doc,
        "The relational schema comprises Block (campus wing identifier), FloorPlan (foreign key to Block, "
        "dimensions, JSON blocked-cell list, optional ImageField), Scan (raw tuples with mode m ∈ {wifi, mobile}, "
        "signal s ∈ ℤ in dBm, timestamps), and CellAggregate (median, scan count, provider key, boolean "
        "all-provider flag). Uniqueness is enforced on (floor, x, y, m, provider_key, all_flag).",
    )

    ieee_heading(doc, "V. AGGREGATION AND INTERPOLATION METHODS")
    body_p(
        doc,
        "For each logical bucket—either a named provider or the composite all-provider view—the aggregate signal "
        "ŝ(x, y) is the sample median of all Scan observations in that bucket sharing (x, y) and m. The median "
        "minimizes L1 deviation in one dimension and is less sensitive than the arithmetic mean to impulsive "
        "RSSI spikes [5]. Even-cardinality tie-breaking follows Python statistics.median (mean of inner two order "
        "statistics). On each insert, refresh executes in one database transaction: re-query signals for the "
        "provider-local bucket and for the union over providers, then upsert both aggregates.",
    )
    body_p(
        doc,
        "Interpolation applies only to cells not in B_ω and lacking measured points. Let P be the set of "
        "grid locations with measured pairs (ŝ_i, n_i) where n_i is the scan count used as a confidence scalar. "
        "Within Chebyshev radius d_max = 2 (excluding the center), for each offset (δx, δy) ≠ (0, 0) with "
        "Euclidean square distance r² = δx² + δy² > 0, neighbor weight is:",
    )
    add_eq_label(doc, "w_i = (1 / r²) · max(1, √n_i)", 2)
    body_p(
        doc,
        "The interpolated value is the weighted average ∑ w_i ŝ_i / ∑ w_i, rounded to two decimals; outputs "
        "carry the flag interpolated = true and count = 0. A single pass avoids propagating estimates "
        "through freshly synthesized cells, limiting false confidence accumulation—a pragmatic design choice "
        "for visualization rather than sequential Kriging [6].",
    )
    body_p(
        doc,
        "Time complexity: registry construction is O(F) over F floor plans; per-scan refresh is O(S_cell) "
        "for scans collocated in the cell; full-floor rebuild is O(S); interpolation is O(N_x N_y · |𝒩|) with "
        "|𝒩| bounded by the fixed window size.",
    )

    ieee_heading(doc, "VI. API AND CLIENT INTERFACES")
    body_p(
        doc,
        "The public read surface exposes GET /api/heatmap/?block=&floor=&mode=&service_provider=&interpolate=. "
        "Provider omission or the literal all selects the precomputed composite bucket. If no aggregate rows "
        "match, the server triggers rebuild_aggregates_for_floor as a self-healing path. GET /api/config/ mirrors "
        "the floor registry for native mobile clients. POST /api/scan/ accepts JSON; it is CSRF-exempt to support "
        "non-browser agents and therefore must be network-restricted in production deployments.",
    )

    ieee_heading(doc, "VII. VISUALIZATION PIPELINE")
    body_p(
        doc,
        "The browser computes dynamic normalization: letting S_real be measured points only, the color scale "
        "uses min and max over S_real when non-empty, else over all points. Normalized t = clamp((s − s_min)/"
        "max(1, s_max − s_min), 0, 1). Blended mode draws radial gradients with spread tied to sample density; "
        "contour mode quantizes t to ⌊6t + 0.5⌋/6 bands. Measured cells use a red–yellow–green thermal ramp; "
        "interpolated cells use a cool palette and lower base opacity. An optional confidence overlay darkens "
        "cells proportional to relative scan count within the current view.",
    )

    ieee_heading(doc, "VIII. SECURITY AND DEPLOYMENT")
    body_p(
        doc,
        "Browser-based scan submission requires authenticated sessions and CSRF tokens. The read APIs are "
        "unauthenticated; sensitivity of floor plans and coverage maps should inform perimeter controls. "
        "Production templates target Gunicorn behind HTTPS with WhiteNoise for static assets; PostgreSQL is "
        "recommended via DATABASE_URL. Uploaded media may require object storage on ephemeral platform-as-a-service "
        "disks [7].",
    )

    ieee_heading(doc, "IX. LIMITATIONS AND FUTURE WORK")
    body_p(
        doc,
        "NSC does not model vertical coupling between floors, multipath, or antenna patterns. Interpolation is "
        "heuristic IDW on a discrete lattice, not physics-based field estimation. Future extensions include "
        "token-authenticated scan APIs, WebSocket-distributed live updates, and export to GeoPDF or GeoJSON "
        "where georeferencing becomes available.",
    )

    ieee_heading(doc, "X. CONCLUSION")
    body_p(
        doc,
        "We described NetSense Campus, an end-to-end Django and JavaScript system for grid-indexed indoor signal "
        "collection and visualization. Median aggregates, dual-namespace provider storage, and intentionally "
        "conservative interpolation preserve interpretability for network operators. The formalism, algorithms, "
        "and interfaces should assist reproducibility and curriculum integration in wireless systems education.",
    )

    # References
    doc.add_paragraph()
    rh = doc.add_paragraph()
    rhr = rh.add_run("REFERENCES")
    set_run_font(rhr, bold=True)
    refs = [
        "A. Goldsmith, Wireless Communications. Cambridge, U.K., Cambridge Univ. Press, 2005.",
        "C. Ware, Information Visualization: Perception for Design, 3rd ed. Burlington, MA, USA: Morgan Kaufmann, 2012.",
        "Remcom, Wireless InSite RF propagation software, product documentation, 2024. [Online]. Available: https://www.remcom.com/",
        "P. Bahl and V. N. Padmanabhan, “RADAR: An in-building RF-based user location and tracking system,” in Proc. IEEE INFOCOM, 2000, pp. 775–784.",
        "R. J. Serfling, “Generalized quantiles and robust statistical procedures,” SIAM Rev., vol. 26, no. 2, pp. 262–264, 1984.",
        "N. Cressie, Statistics for Spatial Data. New York, NY, USA: Wiley, 1993.",
        "Heroku,“Ephemeral filesystem,” Platform Dev. Center. [Online]. Available: https://devcenter.heroku.com/articles/dynos#ephemeral-filesystem",
        "Django Software Foundation, “Django documentation — release 5.x,” 2024. [Online]. Available: https://docs.djangoproject.com/",
        "E. W. Dijkstra, “Go To statement considered harmful,” Commun. ACM, vol. 11, no. 3, pp. 147–148, Mar. 1968.",
        "IEEE Editorial Style Manual for Authors, IEEE Publishing Operations, Piscataway, NJ, USA, 2021.",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
